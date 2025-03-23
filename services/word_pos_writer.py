from docx import Document
from docx.shared import RGBColor
from database import POS_FOLDER
import os
from openai import OpenAI
import re
import unicodedata
from difflib import ndiff  # 수정: 텍스트 차이점 비교를 위해 추가

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

DOWNLOAD_FOLDER = os.path.join(os.path.dirname(__file__), "..", "download")
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

def reflect_changes_with_llm(std_text, proj_text):
    prompt = (
        f"다음은 LNG선 POS 문서의 기존 문장입니다:\n"
        f"{std_text}\n\n"
        f"아래 프로젝트 사양서의 내용을 반영하되, 전체 문장을 새로 쓰지 말고 "
        f"수치, 단어 등 필요한 부분만 최소한으로 수정해줘:\n"
        f"{proj_text}"
    )

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "너는 LNG선 POS 문서를 작성하는 전문가야."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
        )
        revised = response.choices[0].message.content.strip()
        return revised
    except Exception as e:
        print(f"❌ GPT 호출 실패: {e}")
        return proj_text

def normalize_text(text):
    """텍스트 정규화: 공백, 특수 문자, 대소문자, 유니코드 정규화"""
    text = re.sub(r'\s+', ' ', text)  # 모든 공백 문자 제거
    text = unicodedata.normalize('NFKC', text)  # 유니코드 정규화
    text = re.sub(r'[^\w\s]', '', text)  # 특수 문자 제거 (선택 사항)
    return text.strip().lower()  # 공백 제거, 특수문자 제거, 소문자 변환

def highlight_changes_in_word(doc, original_text, llm_text):
    """원본 텍스트와 LLM 텍스트를 비교하여 변경된 부분을 빨간색 굵은 글씨로 표시"""
    # 문단 단위로 분리
    original_paras = original_text.split("\n")
    llm_paras = llm_text.split("\n")

    # 문서 초기화
    for para in doc.paragraphs:
        para.clear()

    # 문단 단위로 비교 및 스타일 적용
    for orig, llm in zip(original_paras, llm_paras):
        para = doc.add_paragraph()
        # 단어 단위로 비교
        orig_words = orig.split()
        llm_words = llm.split()
        diff = list(ndiff(orig_words, llm_words))

        current_run = None
        for token in diff:
            sign = token[:2]
            word = token[2:]

            if sign == '+ ':  # 추가된 단어 (변경된 부분)
                current_run = para.add_run(word + " ")
                current_run.bold = True
                current_run.font.color.rgb = RGBColor(255, 0, 0)  # 빨간색
            elif sign == '- ':  # 삭제된 단어 (무시)
                continue
            else:  # 변경되지 않은 단어
                current_run = para.add_run(word + " ")

    # LLM 텍스트가 더 길 경우 나머지 문단 추가
    if len(llm_paras) > len(original_paras):
        for extra_para in llm_paras[len(original_paras):]:
            para = doc.add_paragraph()
            run = para.add_run(extra_para)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # 추가된 문단 전체를 빨간색 굵은 글씨로

def update_pos_word_with_differences(ship_type, differences, project_number, pos_filename):
    source_path = os.path.join(POS_FOLDER, ship_type, pos_filename)
    if not os.path.exists(source_path):
        print(f"❌ POS Word 파일이 존재하지 않음: {source_path}")
        return None, []

    doc = Document(source_path)
    change_log = []

    # 1. 전체 텍스트 추출
    full_text = "\n".join([para.text for para in doc.paragraphs])

    # 2. 프롬프트 작성
    prompt = f"다음은 {ship_type} 선박의 표준 POS(선박 인도 사양서) 문서입니다.\n{full_text}\n\n다음 변경 사항을 반영하여 프로젝트 사양에 맞게 수정해주세요.\n"
    for diff in differences:
        prompt += f"- {diff['표준 사양서']}  => {diff['프로젝트 사양서']}\n"
    prompt += "\n수정된 POS 문서를 완전한 형태로 제공해주세요."

    # 3. LLM 호출
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "당신은 최고의 선박 POS 문서 작성 전문가입니다."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
        )
        llm_generated_text = response.choices[0].message.content.strip()
        print("✅ LLM이 생성한 텍스트:", llm_generated_text)

    except Exception as e:
        print(f"❌ LLM 호출 실패: {e}")
        return None, []

    # 4. 변경된 부분을 빨간색 굵은 글씨로 표시
    highlight_changes_in_word(doc, full_text, llm_generated_text)

    # 5. 변경 로그 생성
    for diff in differences:
        change_log.append({
            "표준 사양서": diff["표준 사양서"],
            "프로젝트 사양서": diff["프로젝트 사양서"]
        })

    # 6. 결과 저장
    result_filename = pos_filename.replace("STD", project_number)
    result_path = os.path.join(DOWNLOAD_FOLDER, result_filename)
    doc.save(result_path)
    print(f"✅ Word 문서 저장 완료: {result_path}")
    return result_path, change_log